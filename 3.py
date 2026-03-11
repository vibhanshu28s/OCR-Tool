import os
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from docx.shared import Inches
import logging

# Setup basic logging
logging.basicConfig(level=logging.INFO)


class EngineeringToDoc:
    def __init__(self, pdf_path):
        self.pdf_path = pdf_path
        self.doc = Document()

    def convert(self, output_path):
        # We use fitz (PyMuPDF) to handle the drawing/image snapshots
        pdf_images = fitz.open(self.pdf_path)

        # We use pdfplumber to handle the text and table logic
        with pdfplumber.open(self.pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                logging.info(f"Processing Page {i + 1}...")

                # 1. ADD TEXT 'AS IS'
                text = page.extract_text()
                if text:
                    self.doc.add_paragraph(text)

                # 2. ADD TABLES 'AS IS'
                tables = page.extract_tables()
                for table in tables:
                    if not table: continue
                    docx_table = self.doc.add_table(rows=len(table), cols=len(table[0]))
                    docx_table.style = 'Table Grid'
                    for r_idx, row in enumerate(table):
                        for c_idx, cell in enumerate(row):
                            docx_table.cell(r_idx, c_idx).text = str(cell) if cell else ""

                # 3. ADD ENGINEERING DRAWING AS IMAGE
                # This captures the visual layout (drawings, logos, stamps)
                pix = pdf_images[i].get_pixmap(matrix=fitz.Matrix(2, 2))  # High res
                img_path = f"temp_page_{i}.png"
                pix.save(img_path)

                self.doc.add_heading(f'Drawing/Layout Snapshot - Page {i + 1}', level=2)
                self.doc.add_picture(img_path, width=Inches(6.0))

                # Cleanup temp image
                os.remove(img_path)
                self.doc.add_page_break()

        self.doc.save(output_path)
        pdf_images.close()
        logging.info(f"Success! Saved to {output_path}")


# --- EXECUTION ---
if __name__ == "__main__":
    # 1. Update these paths
    INPUT_PDF = "final_pdf.pdf"
    OUTPUT_DOCX = "Final_Engineering_Report.docx"

    # 2. Run the tool
    converter = EngineeringToDoc(INPUT_PDF)
    converter.convert(OUTPUT_DOCX)