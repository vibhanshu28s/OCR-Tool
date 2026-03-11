import os
import pdfplumber
import pandas as pd
from docx import Document
from docx.shared import Pt
import logging

# 1. Setup Logging for Production Monitoring
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


class EngineeringDocScanner:
    def __init__(self, input_pdf):
        self.input_pdf = input_pdf
        self.doc = Document()
        self._setup_doc_style()

    def _setup_doc_style(self):
        """Sets standard professional font styles for the output DOCX."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)

    def extract_metadata(self, page):
        """Specifically targets project identifiers like USID and Site Numbers."""
        text = page.extract_text()
        metadata = {}
        # Example: Looking for patterns found in engineering titles [cite: 3, 41]
        if "SITE NUMBER" in text:
            metadata['Site Number'] = text.split("SITE NUMBER:")[1].split("\n")[0].strip()
        if "USID#" in text:
            metadata['USID'] = text.split("USID#:")[1].split("\n")[0].strip()
        return metadata

    def process_tables(self, page):
        """Extracts and formats tables like Sheet Indices or Line Tables[cite: 2, 7]."""
        tables = page.extract_tables()
        for table in tables:
            if not table or not any(table): continue

            # Create DOCX table
            docx_table = self.doc.add_table(rows=len(table), cols=len(table[0]))
            docx_table.style = 'Table Grid'

            for i, row in enumerate(table):
                for j, cell in enumerate(row):
                    docx_table.cell(i, j).text = str(cell) if cell else ""

    def run(self, output_docx):
        """Main execution loop for the scanner."""
        if not os.path.exists(self.input_pdf):
            logging.error(f"File not found: {self.input_pdf}")
            return

        logging.info(f"Starting extraction for: {self.input_pdf}")

        try:
            with pdfplumber.open(self.input_pdf) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    logging.info(f"Processing Page {page_num + 1}...")

                    # Add Header for each page
                    self.doc.add_heading(f"Page {page_num + 1} Content", level=1)

                    # 1. Extract and Add Text (General Notes/Compliance)
                    text = page.extract_text()
                    if text:
                        self.doc.add_paragraph(text)

                    # 2. Extract and Add Tables (Schedules/Indices) [cite: 2, 10]
                    self.process_tables(page)

                    self.doc.add_page_break()

            self.doc.save(output_docx)
            logging.info(f"Successfully saved production doc to: {output_docx}")

        except Exception as e:
            logging.error(f"Failed to process PDF: {str(e)}")


# --- PRODUCTION EXECUTION BLOCK ---
if __name__ == "__main__":
    # ASSIGN PDF PATH HERE
    # Ensure 'final_pdf.pdf' is in the same directory as this script
    INPUT_FILE = "final_pdf.pdf"
    OUTPUT_FILE = "Extracted_Engineering_Specs.docx"

    scanner = EngineeringDocScanner(INPUT_FILE)
    scanner.run(OUTPUT_FILE)