import pdfplumber
from docx import Document


def create_ocr_doc(pdf_path, doc_output):
    doc = Document()
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            # 1. Extract Text
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)

            # 2. Extract Tables
            tables = page.extract_tables()
            for table in tables:
                t = doc.add_table(rows=len(table), cols=len(table[0]))
                for i, row in enumerate(table):
                    for j, cell in enumerate(row):
                        t.cell(i, j).text = str(cell) if cell else ""

            # 3. Handle Images/Drawings
            # Note: Complex drawings may require cv2 + pytesseract

    doc.save(doc_output)

create_ocr_doc(pdf_path="test_img.png", doc_output="output.doc")